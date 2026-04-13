"""
Script tạo baocao.docx — Hệ thống Quản lý Điểm NTTU
Chạy: python scripts/generate_docx.py
"""

import io
import re
import sys
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
import numpy as np

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT     = Path(__file__).parent.parent
OUT_FILE = ROOT / "baocao.docx"

# ─── Màu ────────────────────────────────────────────────────────
C_NAVY  = RGBColor(0x0C, 0x1E, 0x3F)
C_BLUE  = RGBColor(0x1F, 0x8F, 0xE4)
C_TEXT  = RGBColor(0x1F, 0x29, 0x37)
C_SUB   = RGBColor(0x64, 0x74, 0x8B)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HEX_HEADER = "0C1E3F"   # navy – table header bg
HEX_GRAY   = "EBF0F5"   # light blue-gray – alternating row
HEX_BLUE   = "1F8FE4"   # blue accent

# ─── Helpers ────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_valign(cell, align="center"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)


def set_row_height(row, height_cm: float):
    tr   = row._tr
    trPr = tr.get_or_add_trPr()
    trH  = OxmlElement("w:trHeight")
    trH.set(qn("w:val"), str(int(height_cm * 567)))   # 1 cm ≈ 567 EMU twips
    trH.set(qn("w:hRule"), "exact")
    trPr.append(trH)


def para_spacing(p, before=0, after=60):
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"),  str(after))
    pPr.append(spacing)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_break_type())


def docx_break_type():
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    return br


def force_page_break(doc):
    p   = doc.add_paragraph()
    run = p.add_run()
    br  = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


# ─── Heading ────────────────────────────────────────────────────

def add_h(doc, text: str, level: int):
    """level: 0=Title, 1=H1, 2=H2, 3=H3"""
    p = doc.add_paragraph()
    para_spacing(p, before=120 if level <= 1 else 80,
                    after=60  if level <= 1 else 40)
    run = p.add_run(text)
    run.bold = True
    if level == 0:
        run.font.size  = Pt(20)
        run.font.color.rgb = C_NAVY
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size  = Pt(15)
        run.font.color.rgb = C_NAVY
        # underline
        rPr = run._r.get_or_add_rPr()
        u   = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    elif level == 2:
        run.font.size  = Pt(13)
        run.font.color.rgb = C_BLUE
    else:
        run.font.size  = Pt(11)
        run.font.color.rgb = C_NAVY
    return p


# ─── Table ──────────────────────────────────────────────────────

def clean(text: str) -> str:
    t = text.strip()
    t = re.sub(r"\*\*(.+?)\*\*", r"\1", t)
    t = re.sub(r"`(.+?)`",       r"\1", t)
    return t


def add_styled_table(doc, headers: list, rows: list, col_widths=None):
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set column widths if provided
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    # Header row
    hdr = table.rows[0]
    set_row_height(hdr, 0.85)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, HEX_HEADER)
        set_cell_valign(cell, "center")
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(clean(h))
        run.bold = True
        run.font.size   = Pt(9)
        run.font.color.rgb = C_WHITE
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_spacing(cell.paragraphs[0], 0, 0)

    # Body rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg  = HEX_GRAY if r_idx % 2 == 1 else "FFFFFF"
        set_row_height(row, 0.75)
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            if bg != "FFFFFF":
                set_cell_bg(cell, bg)
            set_cell_valign(cell, "center")
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(clean(cell_text))
            run.font.size = Pt(9)
            run.font.color.rgb = C_TEXT
            para_spacing(cell.paragraphs[0], 0, 0)

    return table


def parse_md_table(lines):
    header, rows = [], []
    for line in lines:
        if not line.strip().startswith("|"):
            continue
        parts = [p for p in line.strip().strip("|").split("|")]
        if all(re.match(r"^[-: ]+$", p.strip()) for p in parts if p.strip()):
            continue
        if not header:
            header = parts
        else:
            rows.append(parts)
    return header, rows


# ─── Image placeholder box ──────────────────────────────────────

def add_image_placeholder(doc, label: str, height_cm: float = 7.0):
    """Add a grey bordered box to paste a screenshot into."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    # height
    tr   = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trH  = OxmlElement("w:trHeight")
    trH.set(qn("w:val"), str(int(height_cm * 567)))
    trH.set(qn("w:hRule"), "exact")
    trPr.append(trH)

    set_cell_bg(cell, "F3F6FA")
    set_cell_valign(cell, "center")
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(f"[ {label} ]")
    run.font.size  = Pt(10)
    run.font.italic = True
    run.font.color.rgb = C_SUB
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()   # spacing after


# ─── Mindmap image ──────────────────────────────────────────────

def draw_mindmap() -> bytes:
    fig, ax = plt.subplots(figsize=(20, 12))
    ax.set_xlim(0, 20)
    ax.set_ylim(0, 12)
    ax.axis("off")
    fig.patch.set_facecolor("#F0F4F8")

    # Colors
    COL_CENTER  = "#0C1E3F"
    COL_BRANCH  = ["#1F8FE4", "#16A34A", "#D97706", "#DC2626", "#7C3AED"]
    COL_LEAF_BG = "#EBF0F5"
    COL_LEAF_TX = "#1F2937"

    # ── Center node ──────────────────────────────────────────────
    cx, cy = 10, 6
    center_box = FancyBboxPatch((cx-2.2, cy-0.55), 4.4, 1.1,
                                boxstyle="round,pad=0.15",
                                facecolor=COL_CENTER, edgecolor="white",
                                linewidth=2, zorder=5)
    ax.add_patch(center_box)
    ax.text(cx, cy, "HỆ THỐNG\nQUẢN LÝ ĐIỂM NTTU",
            ha="center", va="center", fontsize=10, fontweight="bold",
            color="white", zorder=6, linespacing=1.4)

    # ── Branch definitions ────────────────────────────────────────
    branches = [
        {
            "label": "Auth &\nNgười dùng",
            "pos": (2.8, 9.5),
            "leaves": ["Đăng nhập / JWT", "Phân quyền 3 cấp", "Quản lý tài khoản", "Audit log"],
        },
        {
            "label": "Tìm kiếm &\nLọc dữ liệu",
            "pos": (2.8, 4.5),
            "leaves": ["Tìm sinh viên", "Lọc lớp / khoa", "Lọc năm học", "Phân trang"],
        },
        {
            "label": "Nhập điểm &\nImport",
            "pos": (10, 10.8),
            "leaves": ["Nhập thủ công", "Tính điểm tự động", "Import Excel/CSV", "Xem bảng điểm lớp"],
        },
        {
            "label": "AI Dự đoán &\nCảnh báo",
            "pos": (17.2, 9.5),
            "leaves": ["Dự đoán xếp loại", "Confidence %", "Rủi ro cao/TB/thấp", "Lộ trình học tập"],
        },
        {
            "label": "Admin\nQuản trị",
            "pos": (17.2, 4.5),
            "leaves": ["Quản lý Khoa/Môn", "Quản lý Lớp", "Quản lý Sinh viên", "CTĐT 140 tín"],
        },
    ]

    for idx, branch in enumerate(branches):
        bx, by = branch["pos"]
        color  = COL_BRANCH[idx]

        # Line from center to branch
        ax.annotate("", xy=(bx, by),
                    xytext=(cx, cy),
                    arrowprops=dict(arrowstyle="-",
                                    color=color, lw=2.0,
                                    connectionstyle="arc3,rad=0.1"))

        # Branch box
        bw, bh = 2.6, 0.9
        branch_box = FancyBboxPatch((bx - bw/2, by - bh/2), bw, bh,
                                    boxstyle="round,pad=0.1",
                                    facecolor=color, edgecolor="white",
                                    linewidth=1.5, zorder=4)
        ax.add_patch(branch_box)
        ax.text(bx, by, branch["label"],
                ha="center", va="center", fontsize=8.5, fontweight="bold",
                color="white", zorder=5, linespacing=1.35)

        # Leaves
        leaves = branch["leaves"]
        n = len(leaves)
        # Spread leaves around the branch radially away from center
        dx = bx - cx
        dy = by - cy
        angle_base = np.degrees(np.arctan2(dy, dx))

        spread = 50
        angles = np.linspace(angle_base - spread/2, angle_base + spread/2, n)

        for leaf_label, angle in zip(leaves, angles):
            rad = np.radians(angle)
            lx = bx + 2.4 * np.cos(rad)
            ly = by + 1.1 * np.sin(rad)

            ax.annotate("", xy=(lx, ly),
                        xytext=(bx, by),
                        arrowprops=dict(arrowstyle="-",
                                        color=color, lw=1.2,
                                        connectionstyle="arc3,rad=0.05"))

            lw, lh = 2.2, 0.5
            leaf_box = FancyBboxPatch((lx - lw/2, ly - lh/2), lw, lh,
                                      boxstyle="round,pad=0.08",
                                      facecolor=COL_LEAF_BG,
                                      edgecolor=color, linewidth=1.2, zorder=3)
            ax.add_patch(leaf_box)
            ax.text(lx, ly, leaf_label,
                    ha="center", va="center", fontsize=7.5,
                    color=COL_LEAF_TX, zorder=4)

    plt.title("Mindmap Chức Năng Hệ Thống Quản Lý Điểm NTTU",
              fontsize=13, fontweight="bold", color=COL_CENTER, pad=14)
    plt.tight_layout(pad=0.5)

    buf = io.BytesIO()
    plt.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                facecolor=fig.get_facecolor())
    plt.close(fig)
    buf.seek(0)
    return buf.read()


# ─── Screens catalogue ──────────────────────────────────────────

SCREENS = [
    # (route, name_vi, role, description_vi)
    ("/",                         "Trang chủ (Home)",             "Khách",          "Landing page giới thiệu hệ thống, tính năng AI, công nghệ và hướng dẫn sử dụng theo vai trò."),
    ("/login",                    "Đăng nhập",                    "Tất cả",         "Form đăng nhập split-layout: panel trái logo NTTU + tính năng; panel phải email/mật khẩu + JWT."),
    ("/dashboard",                "Dashboard Tổng quan",          "Teacher, Admin", "4 card thống kê, bảng học sinh rủi ro cao có phân trang, biểu đồ cột phân bố xếp loại A/B/C/F."),
    ("/students",                 "Danh sách Sinh viên",          "Admin",          "Bảng danh sách có tìm kiếm, lọc theo khoa/lớp/trạng thái, phân trang server-side, CRUD qua dialog."),
    ("/students/new",             "Thêm Sinh viên",               "Admin",          "Form tạo mới: họ tên, ngày sinh, giới tính, lớp, thông tin phụ huynh, trạng thái, địa chỉ."),
    ("/students/:id",             "Chi tiết Sinh viên",           "Admin",          "Card thông tin cá nhân + bảng lịch sử điểm theo từng học kỳ + kết quả dự đoán AI mới nhất."),
    ("/students/:id/curriculum",  "Tiến độ CTĐT",                 "Admin, Advisor", "Progress bar tổng tín chỉ đã đạt/cần đạt + danh sách môn học toàn CTĐT với trạng thái pass/fail."),
    ("/classes",                  "Quản lý Lớp học phần",         "Admin, Teacher", "Bảng lớp lọc theo năm học/học kỳ/khoa/từ khóa, dialog CRUD tạo/sửa lớp học phần."),
    ("/classes/:id",              "Chi tiết Lớp học phần",        "Admin, Teacher", "Card thông tin lớp (môn, giáo viên, học kỳ) + bảng danh sách sinh viên đã ghi danh."),
    ("/departments",              "Quản lý Khoa",                 "Admin",          "Row thống kê (khoa/môn/lớp/giáo viên) + lưới card khoa, dialog CRUD thêm/sửa khoa."),
    ("/departments/:id",          "Chi tiết Khoa",                "Admin",          "Card tổng quan + tab Môn học / Lớp học phần / Giáo viên với bảng dữ liệu có dialog inline."),
    ("/subjects",                 "Quản lý Môn học",              "Admin",          "Bảng môn học lọc theo khoa/học kỳ/trạng thái; dialog tạo/sửa với tín chỉ, trọng số, loại môn."),
    ("/majors",                   "Quản lý Ngành học",            "Admin",          "Lưới card ngành lọc theo khoa, hiển thị mã, tên, tổng tín chỉ, thời gian đào tạo."),
    ("/curricula",                "Chương trình Đào tạo",         "Admin",          "Lưới card CTĐT theo ngành/khóa tuyển, tổng tín chỉ, nút xem timeline chi tiết."),
    ("/curricula/:id",            "Timeline CTĐT",                "Admin",          "Lưới theo năm × học kỳ, mỗi ô liệt kê môn học kèm tín chỉ và badge loại môn; nút xuất PDF."),
    ("/grades",                   "Nhập Điểm",                    "Teacher",        "Stepper 3 bước: chọn lớp → chọn sinh viên → nhập TX/GK/TH/TKT với tính điểm realtime."),
    ("/grades/class-sheet",       "Bảng Điểm Lớp",                "Teacher, Admin", "Bảng tổng hợp toàn lớp cuộn ngang: TX1/2/3, GK, TH, TKT, Tổng, GPA4, Xếp loại, Đạt/Rớt."),
    ("/grades/import",            "Import Điểm Excel",            "Teacher",        "Stepper 4 bước: chọn lớp → tải mẫu/upload file → preview lỗi → xác nhận import hàng loạt."),
    ("/predictions",              "Dự đoán AI (Lớp)",             "Teacher, Admin", "Panel trái tìm sinh viên; panel phải badge rủi ro, confidence bar, gợi ý can thiệp, link profile."),
    ("/advisor/students",         "Sinh viên phụ trách",          "Advisor",        "Lưới card sinh viên được phân công: avatar, progress bar tín chỉ, badge rủi ro F/chậm tiến độ."),
    ("/advisor/students/:id",     "Chi tiết SV (Advisor)",        "Advisor",        "Tab: Tổng quan (GPA, rủi ro) / Tiến độ CTĐT / Lộ trình AI (kế hoạch học kỳ theo CTĐT 140 tín)."),
    ("/users",                    "Quản lý Tài khoản",            "Admin",          "Bảng user lọc theo vai trò/khoa; dialog tạo/sửa/khóa, gán khoa, gán sinh viên cố vấn."),
    ("/profile",                  "Hồ sơ Giáo viên",              "Teacher",        "Avatar initials, lưới thông tin cá nhân (khoa, mã, điện thoại, ngày), bảng lớp phụ trách."),
    ("/chat",                     "Chat Nội bộ",                   "Tất cả",         "Sidebar danh sách phòng (tab Tất cả/Chưa đọc/Nhóm, badge unread) + khung chat Socket.IO."),
    ("/news",                     "Tin tức",                      "Tất cả",         "Lưới card bài viết tin tức nội bộ NTTU: tiêu đề, ngày đăng, badge 'Mới'."),
    ("/notifications",            "Thông báo",                    "Tất cả",         "Trang thông báo hệ thống — hiện tại hiển thị empty-state (chưa kết nối dữ liệu thật)."),
]


# ─── Document builder ───────────────────────────────────────────

def build_doc():
    doc = Document()

    # ── Page setup ──────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.0)

    # ── Default style ────────────────────────────────────────────
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    # ════════════════════════════════════════════════════════════
    # TRANG BÌA
    # ════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    para_spacing(p, 0, 0)

    for _ in range(3):
        doc.add_paragraph()

    add_h(doc, "BÁO CÁO ĐỒ ÁN TỐT NGHIỆP", 0)

    p = doc.add_paragraph()
    r = p.add_run("Hệ thống Quản lý Điểm và Dự đoán Kết quả Học tập\nsử dụng Machine Learning")
    r.font.size = Pt(14)
    r.font.color.rgb = C_BLUE
    r.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p, 80, 80)

    for line in [
        "Trường: Đại học Nguyễn Tất Thành",
        "Khoa: Công nghệ Thông tin",
        "Năm học: 2025 – 2026",
    ]:
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.size = Pt(11)
        r.font.color.rgb = C_TEXT
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_spacing(p, 0, 20)

    force_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # CHƯƠNG 2
    # ════════════════════════════════════════════════════════════
    add_h(doc, "CHƯƠNG 2: CHI TIẾT DỰ ÁN", 1)

    # ─── 2.1 Mindmap ────────────────────────────────────────────
    add_h(doc, "2.1 Sơ đồ Mindmap Chức Năng Website", 2)

    p = doc.add_paragraph(
        "Hình dưới đây mô tả tổng quan 5 nhóm chức năng chính của hệ thống "
        "cùng các tính năng thuộc mỗi nhóm."
    )
    r = p.runs[0]; r.font.size = Pt(11); r.font.color.rgb = C_TEXT
    para_spacing(p, 0, 60)

    print("  Đang vẽ mindmap...")
    img_bytes = draw_mindmap()
    buf = io.BytesIO(img_bytes)
    doc.add_picture(buf, width=Inches(6.2))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(last_para, 0, 80)

    # ─── 2.2 Chức năng ──────────────────────────────────────────
    add_h(doc, "2.2 Các Chức Năng Chính", 2)

    # 2.2.1
    add_h(doc, "2.2.1 Chức năng Authentication & Người dùng", 3)
    add_styled_table(doc,
        ["Chức năng", "Mô tả", "Vai trò"],
        [
            ["Đăng nhập",          "Xác thực email + mật khẩu, trả về JWT token",             "Tất cả"],
            ["Phân quyền 3 cấp",   "Admin / Giáo viên (Teacher) / Cố vấn (Advisor)",           "Hệ thống"],
            ["Quản lý hồ sơ",      "Xem/sửa thông tin cá nhân, đổi mật khẩu",                  "Tất cả"],
            ["Quản lý tài khoản",  "Admin tạo/sửa/khóa/xóa tài khoản giáo viên & advisor",     "Admin"],
            ["Remember me",        "Lưu phiên đăng nhập 7 ngày qua localStorage",               "Tất cả"],
            ["Bảo vệ route",       "AuthGuard + AdminOnlyGuard kiểm tra token trước route",      "Hệ thống"],
            ["Audit log",          "Ghi lại mọi hành động nhạy cảm vào collection auditlogs",   "Hệ thống"],
        ],
        col_widths=[4.5, 9.0, 3.0]
    )
    doc.add_paragraph()

    p = doc.add_paragraph(); r = p.add_run("Luồng đăng nhập:"); r.bold = True; r.font.size = Pt(10)
    for step in [
        "1. Người dùng nhập email + mật khẩu",
        "2. Backend verify bcrypt + kiểm tra isActive",
        "3. Trả JWT + thông tin user (role, departmentIds)",
        "4. Angular lưu token vào localStorage",
        "5. JWT Interceptor đính kèm Bearer vào mọi request",
        "6. Token hết hạn → redirect /login",
    ]:
        p = doc.add_paragraph(step, style="List Number")
        p.runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # 2.2.2
    add_h(doc, "2.2.2 Chức năng Tìm kiếm & Lọc dữ liệu", 3)
    add_styled_table(doc,
        ["Chức năng", "Mô tả", "Trang áp dụng"],
        [
            ["Tìm sinh viên",            "Tìm theo tên / mã SV / email",                         "Danh sách sinh viên"],
            ["Lọc theo lớp học phần",    "Dropdown chọn lớp → load danh sách",                    "Nhập điểm, Xem điểm"],
            ["Lọc theo khoa",            "Admin lọc sinh viên / lớp theo từng khoa",               "Quản lý"],
            ["Lọc theo năm học / HK",    "Chọn năm học để xem dữ liệu điểm tương ứng",            "Dashboard, Điểm"],
            ["Lọc cảnh báo rủi ro",      "Chỉ hiển thị học sinh rủi ro cao / trung bình",          "Dashboard"],
            ["Tìm kiếm môn học",         "Tìm theo tên / mã môn / tín chỉ",                       "Quản lý môn học"],
            ["Phân trang",               "Tất cả danh sách có paginator Items per page",            "Toàn hệ thống"],
        ],
        col_widths=[4.5, 8.0, 4.0]
    )
    doc.add_paragraph()

    # 2.2.3
    add_h(doc, "2.2.3 Chức năng Nhập điểm & Import hàng loạt", 3)
    add_styled_table(doc,
        ["Chức năng", "Mô tả"],
        [
            ["Nhập điểm thủ công",  "Giáo viên chọn lớp → nhập TX1/TX2/TX3, GK, TH, TKT từng sinh viên"],
            ["Tính điểm tự động",   "finalScore = txAvg×0.3 + GK×0.1 + TH×0.2 + TKT×0.4"],
            ["Quy đổi tự động",     "Tính GPA thang 4, letterGrade (A/B+/…/F) và xếp loại"],
            ["Import Excel/CSV",    "Tải mẫu → điền điểm → upload → preview → import vào DB"],
            ["Validation import",   "Kiểm tra mã SV tồn tại, điểm hợp lệ (0–10) trước khi lưu"],
            ["Prefill template",    "File Excel mẫu đã có sẵn toàn bộ sinh viên của lớp"],
            ["Xem bảng điểm lớp",   "Bảng tổng hợp TX/GK/TH/TKT/Tổng/GPA4/Xếp loại cuộn ngang"],
            ["Gọi AI sau import",   "Sau khi lưu → tự động dự đoán AI hàng loạt cho cả lớp"],
        ],
        col_widths=[5.0, 11.5]
    )
    doc.add_paragraph()

    # 2.2.4
    add_h(doc, "2.2.4 Chức năng Dự đoán AI & Cảnh báo", 3)
    add_styled_table(doc,
        ["Chức năng", "Mô tả"],
        [
            ["Dự đoán xếp loại",    "Random Forest dự đoán: Giỏi / Khá / Trung bình / Yếu"],
            ["Mức độ tự tin",       "Confidence % của mô hình cho từng dự đoán"],
            ["Đánh giá rủi ro",     "Phân loại high / medium / low theo xếp loại + confidence"],
            ["Cảnh báo dashboard",  "Badge đỏ hiển thị tổng sinh viên rủi ro cao trên menu"],
            ["Danh sách cảnh báo",  "Bảng học sinh rủi ro cao phân trang, lọc theo lớp"],
            ["Lộ trình AI",         "Gợi ý kế hoạch học tập theo CTĐT 140 tín theo học kỳ"],
            ["Dự đoán hàng loạt",   "Chạy AI toàn bộ sinh viên lớp sau khi nhập/import điểm"],
            ["Lịch sử dự đoán",     "Xem lịch sử các lần dự đoán theo từng sinh viên"],
        ],
        col_widths=[5.0, 11.5]
    )
    p = doc.add_paragraph()
    p.add_run("Thông tin mô hình AI: ").bold = True
    doc.add_paragraph("- Thuật toán: Random Forest Classifier (scikit-learn)")
    doc.add_paragraph("- Features: 27 điểm môn học (subject_codes)")
    doc.add_paragraph("- Training data: ~20,000 mẫu sinh viên thực tế")
    doc.add_paragraph("- Accuracy: ~85–90% trên tập test")
    doc.add_paragraph("- Serving: FastAPI (Python) port 5000, Node.js gọi qua HTTP")
    doc.add_paragraph()

    # 2.2.5
    add_h(doc, "2.2.5 Chức năng Admin Quản trị", 3)
    add_styled_table(doc,
        ["Chức năng", "Mô tả"],
        [
            ["Quản lý Khoa",            "CRUD khoa đào tạo, thống kê số môn/lớp/giáo viên"],
            ["Quản lý Môn học",         "CRUD môn: tên, mã, số TC, khoa, học kỳ, trọng số"],
            ["Quản lý Lớp học phần",    "CRUD lớp: tên, môn học, giáo viên, năm học, sĩ số"],
            ["Quản lý Năm học",         "CRUD năm học, đánh dấu năm học hiện tại (isCurrent)"],
            ["Quản lý Sinh viên",       "CRUD sinh viên, gán lớp, xem lịch sử điểm"],
            ["Quản lý Tài khoản",       "Tạo/sửa/khóa/xóa tài khoản giáo viên và cố vấn"],
            ["Chương trình đào tạo",    "Quản lý CTĐT 140 tín, gán môn từng học kỳ"],
            ["Dashboard tổng quan",     "Thống kê toàn trường: SV, phân bố xếp loại, cảnh báo"],
        ],
        col_widths=[5.0, 11.5]
    )
    doc.add_paragraph()

    force_page_break(doc)

    # ─── 2.3 Giao diện ──────────────────────────────────────────
    add_h(doc, "2.3 Giao Diện Chính", 2)

    p = doc.add_paragraph(
        "Hệ thống bao gồm 26 trang giao diện. Dưới đây là danh sách đầy đủ "
        "kèm mô tả và ảnh chụp màn hình."
    )
    p.runs[0].font.size = Pt(11)
    para_spacing(p, 0, 60)

    # Overview table
    add_styled_table(doc,
        ["#", "Đường dẫn", "Tên trang", "Vai trò"],
        [[str(i+1), s[0], s[1], s[3][:48] + ("…" if len(s[3]) > 48 else "")]
         for i, s in enumerate(SCREENS)],
        col_widths=[0.7, 4.8, 4.8, 6.2]
    )

    force_page_break(doc)

    # Individual screen detail + placeholder
    add_h(doc, "Chi tiết từng màn hình", 3)
    for i, (route, name_vi, role, desc_vi) in enumerate(SCREENS):
        add_h(doc, f"{i+1}. {name_vi}", 3)

        info_rows = [
            ["Đường dẫn", route],
            ["Vai trò", role],
            ["Mô tả", desc_vi],
        ]
        t = doc.add_table(rows=len(info_rows), cols=2)
        t.style = "Table Grid"
        COL_W = [2.8, 13.7]
        for r_idx, (k, v) in enumerate(info_rows):
            row = t.rows[r_idx]
            c0, c1 = row.cells[0], row.cells[1]
            c0.width = Cm(COL_W[0]); c1.width = Cm(COL_W[1])
            set_cell_bg(c0, "EBF0F5")
            c0.paragraphs[0].clear(); c1.paragraphs[0].clear()
            r0 = c0.paragraphs[0].add_run(k)
            r0.bold = True; r0.font.size = Pt(9); r0.font.color.rgb = C_NAVY
            r1 = c1.paragraphs[0].add_run(v)
            r1.font.size = Pt(9); r1.font.color.rgb = C_TEXT
            para_spacing(c0.paragraphs[0], 0, 0)
            para_spacing(c1.paragraphs[0], 0, 0)
        doc.add_paragraph()

        add_image_placeholder(doc, f"Ảnh chụp màn hình: {name_vi}", height_cm=7.0)

        if i < len(SCREENS) - 1:
            force_page_break(doc)

    force_page_break(doc)

    # ─── 2.4 Công nghệ ──────────────────────────────────────────
    add_h(doc, "2.4 Công Nghệ Frontend & Backend", 2)
    add_h(doc, "Bảng công nghệ", 3)

    add_styled_table(doc,
        ["Tầng", "Công nghệ", "Phiên bản", "Mục đích"],
        [
            ["Frontend",    "Angular",          "21.2",    "SPA framework, Standalone Components"],
            ["",            "Angular Material", "21.x MDC","UI components (form, table, dialog)"],
            ["",            "Chart.js",         "4.x",     "Biểu đồ cột, donut"],
            ["",            "Lucide Angular",   "Latest",  "Icon library"],
            ["",            "Socket.IO Client", "4.x",     "Chat realtime"],
            ["Backend",     "Node.js",          "20.x",    "Runtime"],
            ["",            "Express.js",       "4.x",     "HTTP framework, routing"],
            ["",            "Mongoose",         "8.x",     "ODM cho MongoDB"],
            ["",            "bcryptjs",         "2.x",     "Hash mật khẩu"],
            ["",            "jsonwebtoken",     "9.x",     "JWT auth"],
            ["",            "multer + xlsx",    "Latest",  "Upload & parse file Excel"],
            ["",            "Socket.IO",        "4.x",     "Chat realtime server"],
            ["",            "axios",            "1.x",     "Gọi Python AI Engine"],
            ["AI Engine",   "Python",           "3.11",    "Runtime"],
            ["",            "FastAPI",          "0.110+",  "HTTP API server cho model"],
            ["",            "scikit-learn",     "1.4+",    "Random Forest Classifier"],
            ["",            "pandas / numpy",   "Latest",  "Xử lý dữ liệu training"],
            ["",            "pymongo",          "4.x",     "Đọc dữ liệu từ MongoDB để train"],
            ["Database",    "MongoDB Atlas",    "7.x",     "NoSQL cloud database"],
            ["Auth",        "JWT (HS256)",       "—",       "Stateless authentication"],
            ["DevOps",      "concurrently",     "Latest",  "Chạy 3 tiến trình song song"],
        ],
        col_widths=[2.5, 3.5, 2.5, 8.0]
    )
    doc.add_paragraph()

    add_h(doc, "Cấu trúc folder", 3)
    folder_text = (
        "NguyenTatThanhGradeManager/\n"
        "├── backend/src/\n"
        "│   ├── config/database.js        # Kết nối MongoDB Atlas\n"
        "│   ├── middleware/               # auth, adminOnly, advisorAccess, auditLog\n"
        "│   ├── models/                  # 13 Mongoose schemas\n"
        "│   ├── routes/                  # 13 Express routers\n"
        "│   ├── services/                # aiService, importService, curriculumService\n"
        "│   └── index.js\n"
        "├── ai-engine/\n"
        "│   ├── data/generate_data.py    # Sinh dataset từ MongoDB\n"
        "│   ├── models/                  # model.pkl, feature_names.pkl\n"
        "│   ├── main.py                  # FastAPI /predict /health\n"
        "│   ├── train.py                 # Train Random Forest\n"
        "│   └── risk_logic.py            # Tính riskLevel\n"
        "└── frontend/src/app/\n"
        "    ├── core/                    # guards, interceptors, services\n"
        "    ├── features/                # 16 lazy-loaded feature modules\n"
        "    └── shared/                  # layout, models, styles"
    )
    p = doc.add_paragraph(folder_text)
    p.runs[0].font.name = "Courier New"
    p.runs[0].font.size = Pt(8.5)
    p.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0x6B)

    force_page_break(doc)

    # ─── 2.5 Database ───────────────────────────────────────────
    add_h(doc, "2.5 Database Models", 2)
    add_h(doc, "Bảng các Collections", 3)

    add_styled_table(doc,
        ["#", "Collection", "Mô tả", "Trường chính"],
        [
            ["1",  "users",             "Tài khoản hệ thống",          "email, passwordHash, role, departmentIds[], isActive"],
            ["2",  "departments",       "Khoa đào tạo",                "name, code, description, isActive"],
            ["3",  "schoolyears",       "Năm học",                     "name, startYear, endYear, isCurrent"],
            ["4",  "subjects",          "Môn học",                     "name, code, credits, departmentId, semester, weights"],
            ["5",  "majors",            "Ngành học",                   "name, code, departmentId, totalCredits"],
            ["6",  "curricula",         "Chương trình đào tạo",        "majorId, schoolYearId, totalCredits, semesters[]"],
            ["7",  "classes",           "Lớp học phần",                "name, subjectId, teacherId, schoolYearId, studentCount"],
            ["8",  "students",          "Sinh viên",                   "studentId, fullName, email, classId, majorId, cohort"],
            ["9",  "studentcurricula",  "Đăng ký CTĐT của SV",        "studentId, curriculumId, registrations[]"],
            ["10", "grades",            "Điểm theo lớp học phần",      "studentId, classId, tx[], gk, th, tkt, finalScore, gpa4, letterGrade"],
            ["11", "predictions",       "Kết quả dự đoán AI",          "studentId, gradeId, predictedGrade, confidence, riskLevel"],
            ["12", "messages",          "Tin nhắn chat",               "roomId, senderId, content, type, createdAt"],
            ["13", "auditlogs",         "Nhật ký hành động",           "userId, action, targetModel, targetId, changes, ip"],
        ],
        col_widths=[0.6, 3.0, 3.8, 9.1]
    )
    doc.add_paragraph()

    p = doc.add_paragraph(); r = p.add_run("Quan hệ giữa các collections:"); r.bold = True; r.font.size = Pt(10)
    rel_text = (
        "Department ──< Subject ──< Class ──< Grade ──< Prediction\n"
        "                                       │\n"
        "Student ─────────────────────────────<┘\n"
        "   │\n"
        "   └──< StudentCurriculum ──> Curriculum ──> Major ──> Department"
    )
    p2 = doc.add_paragraph(rel_text)
    p2.runs[0].font.name = "Courier New"
    p2.runs[0].font.size = Pt(9)
    p2.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0x6B)

    force_page_break(doc)

    # ─── 2.6 Hướng phát triển ───────────────────────────────────
    add_h(doc, "2.6 Hướng Phát Triển", 2)

    add_h(doc, "Ngắn hạn (1–3 tháng)", 3)
    add_styled_table(doc,
        ["#", "Tính năng", "Mô tả"],
        [
            ["1", "Thông báo realtime",      "Push notification khi cảnh báo rủi ro mới (Socket.IO alert:new)"],
            ["2", "Xuất báo cáo PDF",        "Xuất bảng điểm / báo cáo AI bằng jspdf + html2canvas"],
            ["3", "Dashboard Advisor",       "Dashboard riêng cho cố vấn: thống kê SV phụ trách, tỷ lệ rủi ro"],
            ["4", "Lịch sử điểm sinh viên",  "Timeline GPA qua nhiều học kỳ, biểu đồ tiến trình"],
            ["5", "Reset mật khẩu",          "Quên mật khẩu qua email (Nodemailer + OTP 6 số)"],
        ],
        col_widths=[0.6, 4.5, 11.4]
    )
    doc.add_paragraph()

    add_h(doc, "Trung hạn (3–6 tháng)", 3)
    add_styled_table(doc,
        ["#", "Tính năng", "Mô tả"],
        [
            ["1", "Retrain model tự động",  "Cron job hàng tháng generate → train → deploy model mới"],
            ["2", "Multi-model AI",         "Thử nghiệm XGBoost, Neural Network so sánh với Random Forest"],
            ["3", "Mobile App",             "Flutter app cho phụ huynh theo dõi điểm / cảnh báo"],
            ["4", "Tích hợp LMS",           "Đồng bộ điểm với Moodle / Canvas qua REST API"],
            ["5", "API công khai",          "Cổng tra cứu điểm cho sinh viên qua mã SV + OTP"],
        ],
        col_widths=[0.6, 4.5, 11.4]
    )
    doc.add_paragraph()

    add_h(doc, "Dài hạn (6–12 tháng)", 3)
    add_styled_table(doc,
        ["#", "Tính năng", "Mô tả"],
        [
            ["1", "Phân tích toàn trường",   "Báo cáo tổng hợp theo khoa/ngành/năm, tỷ lệ đạt/rớt"],
            ["2", "Recommendation engine",   "Gợi ý môn học phù hợp, thứ tự đăng ký theo năng lực SV"],
            ["3", "Mở rộng trường khác",     "Multi-tenant hỗ trợ nhiều trường trên cùng nền tảng"],
            ["4", "Phân tích văn bản",       "NLP phân tích nhận xét GV, khảo sát SV bổ sung features AI"],
            ["5", "Tích hợp SSO",            "Đăng nhập một lần qua Microsoft 365 / Google Workspace"],
        ],
        col_widths=[0.6, 4.5, 11.4]
    )

    # ─── Save ────────────────────────────────────────────────────
    doc.save(OUT_FILE)
    print(f"✅ Đã tạo: {OUT_FILE}")


if __name__ == "__main__":
    build_doc()
